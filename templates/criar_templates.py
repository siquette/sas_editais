"""
templates/criar_templates.py
Gera os dois templates do pacote de Contratação Direta por Dispensa (art. 75, II):
  - tr_fornecimento.docx
  - aviso_contratacao.docx

Baseado nos modelos pré-aprovados pela PGU-USP (Portaria PG nº 12/24).
Os templates são gerados automaticamente na primeira execução do app.
"""

from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Helpers ────────────────────────────────────────────────────────────────

def _r(par, texto, bold=False, italic=False, pt=11):
    run = par.add_run(texto)
    run.bold, run.italic = bold, italic
    run.font.size = Pt(pt)
    return run

def _p(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    return p

def _titulo(doc, texto, pt=13):
    p = _p(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _r(p, texto, bold=True, pt=pt)
    return p

def _secao(doc, num, texto):
    p = _p(doc, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(10)
    _r(p, f"{num}. {texto.upper()}", bold=True, pt=11)
    return p

def _subsecao(doc, num, texto):
    p = _p(doc, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(6)
    _r(p, f"{num} ", bold=True, pt=11)
    _r(p, texto, bold=True, italic=True, pt=11)
    return p

def _assinatura(doc, rotulo):
    p = _p(doc, WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_before = Pt(28)
    _r(p, "_" * 52)
    p2 = _p(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _r(p2, rotulo, pt=10)


# ══════════════════════════════════════════════════════════════════════════
# TEMPLATE 1 — TERMO DE REFERÊNCIA PARA FORNECIMENTO
# Art. 6º, XXIII; art. 40, §1º; art. 72 — Lei 14.133/2021
# Modelo: TR 2.1 pré-aprovado PGU-USP (13/03/2024)
# ══════════════════════════════════════════════════════════════════════════

def criar_tr(destino: Path):
    from config import settings
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(2)

    # Cabeçalho
    _titulo(doc, "{{ instituicao_nome | upper }}", pt=12)
    _titulo(doc, "{{ instituicao_setor | upper }}", pt=11)
    doc.add_paragraph()
    _titulo(doc, "TERMO DE REFERÊNCIA", pt=14)
    _titulo(doc, "CONTRATAÇÃO DIRETA — DISPENSA POR VALOR", pt=11)
    p = _p(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _r(p, "{{ fundamento_legal }}", italic=True, pt=10)
    doc.add_paragraph()

    # Tabela de identificação
    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = "Table Grid"
    dados = [
        ("Processo SEI nº", "{{ numero_sei or 'A preencher' }}",
         "Data",            "{{ data_geracao }}"),
        ("Unidade Solicitante", "{{ unidade_solicitante }}",
         "Agente de Contratação", "{{ agente_contratacao or 'A designar' }}"),
        ("Fundamento Legal", "{{ fundamento_legal }}",
         "Inciso", "{{ inciso_legal or '—' }}"),
    ]
    for i, (r1, v1, r2, v2) in enumerate(dados):
        row = tbl.rows[i]
        for j, txt in enumerate([r1, v1, r2, v2]):
            c = row.cells[j]
            c.text = txt
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
                run.bold = (j % 2 == 0)
    doc.add_paragraph()

    # 1. OBJETO
    _secao(doc, "1", "DO OBJETO")
    p = _p(doc); _r(p, "1.1 Objeto: "); _r(p, "{{ objeto }}", bold=True); _r(p, ".")
    p = _p(doc); _r(p, "1.2 ")
    _r(p, "{%- if tem_catmat %}Código CATMAT/CATSER: {{ codigo_catmat }}.{%- endif %}", italic=True, pt=10)
    p = _p(doc)
    _r(p, "1.3 Quantidade: "); _r(p, "{{ quantidade }} {{ unidade_medida }}", bold=True); _r(p, ".")

    # 2. JUSTIFICATIVA
    _secao(doc, "2", "DA JUSTIFICATIVA E FUNDAMENTAÇÃO DA NECESSIDADE")
    _subsecao(doc, "2.1", "Da necessidade")
    _p(doc); doc.paragraphs[-1].add_run("{{ justificativa_necessidade }}").font.size = Pt(11)

    _subsecao(doc, "2.2", "Da justificativa da quantidade")
    p = _p(doc)
    _r(p, "{{ justificativa_quantidade or 'A quantidade solicitada foi definida com base na demanda prevista para o período, considerando o uso regular pela unidade solicitante.' }}")

    _subsecao(doc, "2.3", "Do enquadramento legal")
    p = _p(doc)
    _r(p, "A presente contratação enquadra-se na hipótese do ")
    _r(p, "{{ fundamento_legal }}", bold=True)
    _r(p, ", haja vista que o valor estimado de ")
    _r(p, "{{ valor_total_estimado }}", bold=True)
    _r(p, " é inferior ao limite legal vigente (Decreto nº 12.343/2024). "
         "Respeita-se o somatório das contratações de mesma natureza no "
         "exercício financeiro, nos termos do art. 75, § 1º, da Lei nº 14.133/2021.")

    # 3. ESPECIFICAÇÕES TÉCNICAS
    _secao(doc, "3", "DA DESCRIÇÃO DO OBJETO E ESPECIFICAÇÕES TÉCNICAS")
    _subsecao(doc, "3.1", "Especificações mínimas")
    p = _p(doc); _r(p, "O objeto deverá atender às seguintes especificações técnicas mínimas:")
    p = _p(doc); _r(p, "{{ especificacoes_lista }}")

    _subsecao(doc, "3.2", "Critério de aceitação")
    p = _p(doc)
    _r(p, "{{ criterio_aceitacao or 'O recebimento definitivo fica condicionado à verificação "
         "da conformidade do bem com as especificações constantes neste Termo de Referência, "
         "a ser realizada pelo fiscal designado no prazo de 5 (cinco) dias úteis após "
         "o recebimento provisório.' }}")

    _subsecao(doc, "3.3", "Garantia e assistência técnica")
    p = _p(doc)
    _r(p, "{%- if tem_garantia %}Garantia: {{ garantia_produto }}. "
         "{%- else %}O fornecedor deverá garantir o bem pelo prazo mínimo estabelecido "
         "pelo fabricante, assegurando assistência técnica durante o período de garantia."
         "{%- endif %}")

    # 4. MODELO DE EXECUÇÃO
    _secao(doc, "4", "DO MODELO DE EXECUÇÃO DO OBJETO")
    _subsecao(doc, "4.1", "Do prazo e local de entrega")
    p = _p(doc)
    _r(p, "O bem deverá ser entregue no prazo de ")
    _r(p, "{{ prazo_entrega_dias }} dias corridos", bold=True)
    _r(p, ", contados da emissão da Ordem de Fornecimento, no endereço: ")
    _r(p, "{{ local_entrega or 'a ser indicado pela unidade solicitante.' }}")

    _subsecao(doc, "4.2", "Das condições de entrega")
    p = _p(doc)
    _r(p, "{{ condicoes_entrega or 'O bem deverá ser entregue em embalagem original do "
         "fabricante, sem avarias, acompanhado de nota fiscal, manual de instruções "
         "(quando aplicável) e demais documentos necessários.' }}")

    _subsecao(doc, "4.3", "Do recebimento")
    p = _p(doc)
    _r(p, "4.3.1 Recebimento "); _r(p, "provisório:", italic=True)
    _r(p, " realizado pelo servidor designado, no ato da entrega, para verificação "
         "da conformidade com as especificações e quantidades contratadas.")
    p = _p(doc)
    _r(p, "4.3.2 Recebimento "); _r(p, "definitivo:", italic=True)
    _r(p, " realizado pelo fiscal em até 5 (cinco) dias úteis após o recebimento "
         "provisório, mediante atesto na Nota Fiscal.")

    # 5. MODELO DE GESTÃO
    _secao(doc, "5", "DO MODELO DE GESTÃO DO CONTRATO")
    _subsecao(doc, "5.1", "Da fiscalização")
    p = _p(doc)
    _r(p, "A execução será fiscalizada por ")
    _r(p, "{{ fiscal_contrato or 'servidor(a) a ser designado(a) pela autoridade competente' }}")
    _r(p, ", nos termos do art. 117 da Lei nº 14.133/2021.")

    _subsecao(doc, "5.2", "Obrigações do fornecedor")
    obrig_forn = [
        "Entregar o objeto nas condições, prazos e locais estabelecidos neste TR;",
        "Substituir, no prazo de 5 (cinco) dias úteis, o bem recusado ou entregue em "
        "desacordo com as especificações;",
        "Emitir Nota Fiscal correspondente ao objeto entregue;",
        "Manter, durante toda a execução, as condições de habilitação exigidas.",
    ]
    for o in obrig_forn:
        p = doc.add_paragraph(o, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    _subsecao(doc, "5.3", "Obrigações da contratante")
    obrig_cont = [
        "Proporcionar todas as condições para o recebimento do objeto;",
        "Efetuar o pagamento nas condições e prazos estabelecidos;",
        "Notificar o fornecedor sobre irregularidades verificadas.",
    ]
    for o in obrig_cont:
        p = doc.add_paragraph(o, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    # 6. HABILITAÇÃO
    _secao(doc, "6", "DOS CRITÉRIOS DE SELEÇÃO DO FORNECEDOR")
    _subsecao(doc, "6.1", "Da habilitação")
    p = _p(doc)
    _r(p, "Para habilitação, verificar a regularidade do fornecedor nos seguintes "
         "cadastros (art. 72, inciso V, da Lei nº 14.133/2021 e Decreto Estadual "
         "nº 68.304/2024):")
    cadastros = [
        "SICAF — Sistema de Cadastro Unificado de Fornecedores;",
        "CEIS — Cadastro Nacional de Empresas Inidôneas e Suspensas (CGU);",
        "CNEP — Cadastro Nacional de Empresas Punidas (CGU);",
        "CNCIAI — Cadastro Nacional de Condenações por Improbidade Administrativa (CNJ);",
        "e-Sanções — Sistema Eletrônico de Aplicação e Registro de Sanções (SP);",
        "CEEP — Cadastro Estadual de Empresas Punidas (SP);",
        "Apenados TCESP.",
    ]
    for c in cadastros:
        p = doc.add_paragraph(c, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    p = _p(doc)
    _r(p, "Obs.: ", bold=True)
    _r(p, "Para contratações com entrega imediata (até 30 dias) ou valor inferior a "
         "{{ limite_habilitacao }} (¼ do limite de dispensa), exige-se apenas regularidade "
         "perante a Fazenda Estadual, Justiça do Trabalho e Seguridade Social "
         "(art. 18 do Decreto nº 68.304/2024).", pt=10)

    # 7. ESTIMATIVA DE PREÇOS
    _secao(doc, "7", "DA ESTIMATIVA DE PREÇOS E DO VALOR ESTIMADO")
    p = _p(doc)
    _r(p, "7.1 O valor estimado, apurado por pesquisa de preços (Decreto Estadual "
         "nº 67.888/2023), é de:")
    p = _p(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _r(p, "{%- if tem_valor_unitario %}Valor unitário estimado: {{ valor_unitario_estimado }}   {%- endif %}")
    p = _p(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _r(p, "Valor total estimado: ", bold=True, pt=12)
    _r(p, "{{ valor_total_estimado }}", bold=True, pt=12)
    p = _p(doc)
    _r(p, "7.2 A memória de cálculo e os documentos que embasam a estimativa "
         "constam do processo administrativo.")

    # 8. DOTAÇÃO
    _secao(doc, "8", "DA ADEQUAÇÃO ORÇAMENTÁRIA")
    p = _p(doc)
    _r(p, "8.1 Dotação Orçamentária: ", bold=True)
    _r(p, "{{ dotacao_orcamentaria or 'Conforme documento de reserva orçamentária anexo.' }}")

    # 9. PAGAMENTO
    _secao(doc, "9", "DAS CONDIÇÕES DE PAGAMENTO")
    p = _p(doc)
    _r(p, "9.1 O pagamento será efetuado em ")
    _r(p, "{{ condicoes_pagamento }}", bold=True)
    _r(p, ".")
    p = _p(doc)
    _r(p, "9.2 O pagamento fica condicionado à regularidade fiscal do fornecedor "
         "e ao atesto da Nota Fiscal pelo fiscal do contrato.")

    # 10. SUSTENTABILIDADE (condicional)
    p = _p(doc); _r(p, "{%- if tem_sustentabilidade %}", pt=1)
    _secao(doc, "10", "DOS REQUISITOS DE SUSTENTABILIDADE")
    p = _p(doc); _r(p, "{{ sustentabilidade }}")
    p = _p(doc); _r(p, "{%- endif %}", pt=1)

    # 11. SANÇÕES
    _secao(doc, "11", "DAS SANÇÕES ADMINISTRATIVAS")
    p = _p(doc)
    _r(p, "O descumprimento das obrigações sujeitará o contratado às sanções dos "
         "arts. 155 a 163 da Lei nº 14.133/2021, garantidos o contraditório e "
         "a ampla defesa, podendo ser aplicadas:")
    sancoes = [
        "Advertência;",
        "Multa de 10% sobre o valor do contrato em caso de inexecução parcial;",
        "Multa de 20% sobre o valor do contrato em caso de inexecução total;",
        "Impedimento de licitar e contratar pelo prazo de até 3 (três) anos;",
        "Declaração de inidoneidade para licitar ou contratar.",
    ]
    for s in sancoes:
        p = doc.add_paragraph(s, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    # 12. DISPOSIÇÕES GERAIS
    _secao(doc, "12", "DAS DISPOSIÇÕES GERAIS")
    p = _p(doc)
    _r(p, "12.1 Os casos omissos serão resolvidos pela autoridade competente, "
         "com observância dos princípios do art. 5º da Lei nº 14.133/2021.")
    p = _p(doc)
    _r(p, "12.2 Este Termo de Referência foi elaborado com base no modelo "
         "padronizado e pré-aprovado pela Procuradoria Geral, nos termos da "
         "Portaria PG nº 12/2024.")

    # 13. OBSERVAÇÕES (condicional)
    p = _p(doc); _r(p, "{%- if tem_observacoes %}", pt=1)
    _secao(doc, "13", "DAS OBSERVAÇÕES ADICIONAIS")
    p = _p(doc); _r(p, "{{ observacoes }}")
    p = _p(doc); _r(p, "{%- endif %}", pt=1)

    # Assinaturas
    doc.add_paragraph()
    p = _p(doc, WD_ALIGN_PARAGRAPH.LEFT)
    _r(p, "{{ instituicao_cidade }}, {{ data_geracao }}")
    doc.add_paragraph()
    _assinatura(doc,
        "Servidor(a) responsável\n"
        "{{ agente_contratacao or 'Nome / Matrícula' }}\n"
        "{{ unidade_solicitante }}")
    doc.add_paragraph()
    _assinatura(doc,
        "Autoridade Competente\n"
        "{{ instituicao_setor }}")

    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destino)
    print(f"  ✓ TR Fornecimento: {destino.name}")


# ══════════════════════════════════════════════════════════════════════════
# TEMPLATE 2 — AVISO DE CONTRATAÇÃO DIRETA
# Art. 72, parágrafo único; art. 54 — Lei 14.133/2021
# Modelo: Aviso 1 pré-aprovado PGU-USP (27/02/2024)
# ══════════════════════════════════════════════════════════════════════════

def criar_aviso(destino: Path):
    from config import settings
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(2)

    _titulo(doc, "{{ instituicao_nome | upper }}", pt=12)
    _titulo(doc, "{{ instituicao_setor | upper }}", pt=11)
    doc.add_paragraph()
    _titulo(doc, "AVISO DE CONTRATAÇÃO DIRETA", pt=14)
    p = _p(doc, WD_ALIGN_PARAGRAPH.CENTER)
    _r(p, "{{ fundamento_legal }} — Dispensa por Valor", italic=True, pt=10)
    doc.add_paragraph()

    # Tabela de identificação
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"
    campos_aviso = [
        ("Processo SEI nº",    "{{ numero_sei or 'A preencher' }}"),
        ("Unidade",            "{{ unidade_solicitante }}"),
        ("Objeto",             "{{ objeto }}"),
        ("Fundamento Legal",   "{{ fundamento_legal }}"),
        ("Valor estimado",     "{{ valor_total_estimado }}"),
        ("Data de divulgação", "{{ data_geracao }}"),
    ]
    for i, (rot, val) in enumerate(campos_aviso):
        row = tbl.rows[i]
        row.cells[0].text = rot
        row.cells[1].text = val
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(10)
    doc.add_paragraph()

    p = _p(doc)
    _r(p, "A "); _r(p, "{{ instituicao_nome }}", bold=True)
    _r(p, " torna público que realizará contratação direta por dispensa de "
         "licitação, nos termos do {{ fundamento_legal }}, conforme especificações "
         "constantes do Termo de Referência.")

    p = _p(doc)
    _r(p, "O prazo de entrega é de ")
    _r(p, "{{ prazo_entrega_dias }} dias corridos", bold=True)
    _r(p, " a partir da Ordem de Fornecimento.")

    p = _p(doc)
    _r(p, "Informações: ", bold=True)
    _r(p, "{{ contato_responsavel }}")

    doc.add_paragraph()
    p = _p(doc, WD_ALIGN_PARAGRAPH.LEFT)
    _r(p, "{{ instituicao_cidade }}, {{ data_geracao }}")
    doc.add_paragraph()
    _assinatura(doc, "Autoridade Competente\n{{ instituicao_setor }}")

    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destino)
    print(f"  ✓ Aviso de Contratação: {destino.name}")


# ── Garantia de existência dos templates ──────────────────────────────────

def garantir_templates():
    """
    Cria os templates se não existirem.
    Chamado automaticamente na inicialização do app.
    """
    from config import settings
    criados = False

    if not settings.TEMPLATE_TR.exists():
        criar_tr(settings.TEMPLATE_TR)
        criados = True

    if not settings.TEMPLATE_AVISO.exists():
        criar_aviso(settings.TEMPLATE_AVISO)
        criados = True

    return criados


if __name__ == "__main__":
    from config import settings
    print("Gerando templates...")
    criar_tr(settings.TEMPLATE_TR)
    criar_aviso(settings.TEMPLATE_AVISO)
    print("Concluído.")
